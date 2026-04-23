"""
Generate Spanish translation deliverables from src/i18n/locales/{en,es}.json.
Outputs to deliverables/:
  - butterfly-challenge-spanish.docx   (side-by-side EN / ES by section)
  - butterfly-challenge-spanish.xlsx   (clean reviewer grid, no technical keys)
"""
import json
import os
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH

from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from openpyxl.worksheet.datavalidation import DataValidation

ROOT = Path(__file__).resolve().parents[1]
EN = json.loads((ROOT / "src/i18n/locales/en.json").read_text(encoding="utf-8"))
ES = json.loads((ROOT / "src/i18n/locales/es.json").read_text(encoding="utf-8"))
OUT = ROOT / "deliverables"
OUT.mkdir(exist_ok=True)

SECTION_TITLES = {
    "nav": "Navigation Bar",
    "home": "Home Page",
    "live": "Live Feed (Home & Live Page)",
    "footer": "Footer",
    "language": "Language Switcher",
    "faq": "FAQ",
    "highlights": "Highlights Carousel",
    "stepTabs": "How It Works — Step Tabs",
    "signBuilder": "Sign Builder Tutorial",
    "chain": "Butterfly Effect Chain",
    "popups": "Popups & Modals",
    "countdown": "Countdown Timer",
    "joinC": "Join the Challenge — Popup",
    "reminderC": "Reminder Popup",
    "support": "Get Support Panel",
    "visualTimeline": "Visual Timeline (Story Page)",
    "shareC": "Share Sheet",
    "livePage": "Live Page",
    "story": "Story Page",
    "science": "Science Page",
    "alliancePage": "Alliance Page",
    "workingProgress": "Working Progress Holding Page",
    "toast": "Toast Messages",
}


def flatten(obj, prefix=""):
    if isinstance(obj, dict):
        for k, v in obj.items():
            child = f"{prefix}.{k}" if prefix else k
            yield from flatten(v, child)
    elif isinstance(obj, list):
        for i, v in enumerate(obj):
            yield from flatten(v, f"{prefix}[{i}]")
    else:
        yield (prefix, obj)


def rows_for_section(section_key):
    en_sec = EN.get(section_key, {})
    es_sec = ES.get(section_key, {})
    en_flat = dict(flatten(en_sec))
    es_flat = dict(flatten(es_sec))
    return [(k, str(en_flat[k]), str(es_flat.get(k, ""))) for k in en_flat]


# ---------- Word (.docx) ----------
def build_docx():
    doc = Document()
    for section in doc.sections:
        section.left_margin = Cm(1.6)
        section.right_margin = Cm(1.6)
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)

    title = doc.add_heading("Butterfly Challenge — Spanish Translation", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Side-by-side English / Español, organised by page section")
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x6E, 0x6E, 0x73)
    doc.add_paragraph()

    for section_key, section_title in SECTION_TITLES.items():
        rows = rows_for_section(section_key)
        if not rows:
            continue
        doc.add_heading(section_title, level=1)
        table = doc.add_table(rows=1, cols=2)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "English"
        hdr[1].text = "Español"
        for cell in hdr:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.bold = True
                    r.font.size = Pt(11)
        for _, en_text, es_text in rows:
            row = table.add_row().cells
            row[0].text = en_text
            row[1].text = es_text
            for cell in row:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(10.5)
        for col in table.columns:
            for cell in col.cells:
                cell.width = Inches(3.2)
        doc.add_paragraph()

    path = OUT / "butterfly-challenge-spanish.docx"
    doc.save(path)
    return path


# ---------- Excel (.xlsx) — clean reviewer sheet ----------
def build_xlsx():
    wb = Workbook()
    ws = wb.active
    ws.title = "Review"

    # --- styles ---
    title_font = Font(bold=True, size=16, color="FFFFFF")
    title_fill = PatternFill("solid", fgColor="0A7B77")
    intro_font = Font(italic=True, size=10, color="6E6E73")

    section_font = Font(bold=True, size=12, color="FFFFFF")
    section_fill = PatternFill("solid", fgColor="0A7B77")

    header_font = Font(bold=True, size=11, color="1D1D1F")
    header_fill = PatternFill("solid", fgColor="E8F5F3")

    thin = Side(border_style="thin", color="D2D2D7")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    wrap = Alignment(wrap_text=True, vertical="top", horizontal="left")
    center = Alignment(wrap_text=True, vertical="center", horizontal="center")

    # --- title row ---
    ws.merge_cells("A1:D1")
    ws["A1"] = "Butterfly Challenge — Spanish Translation Review"
    ws["A1"].font = title_font
    ws["A1"].fill = title_fill
    ws["A1"].alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[1].height = 36

    ws.merge_cells("A2:D2")
    ws["A2"] = "Read each Spanish line. Mark ✓ in Status if it's fine, or write a suggestion in Notes. That's it."
    ws["A2"].font = intro_font
    ws["A2"].alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[2].height = 22

    row = 4  # leave one blank row after intro

    # --- walk each page ---
    for section_key, section_title in SECTION_TITLES.items():
        items = rows_for_section(section_key)
        if not items:
            continue

        # section banner row — merged across all 4 columns
        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=4)
        cell = ws.cell(row=row, column=1, value=section_title)
        cell.font = section_font
        cell.fill = section_fill
        cell.alignment = Alignment(horizontal="left", vertical="center", indent=1)
        ws.row_dimensions[row].height = 26
        row += 1

        # per-section header row
        headers = ["English", "Español", "Status", "Notes"]
        for col, h in enumerate(headers, start=1):
            c = ws.cell(row=row, column=col, value=h)
            c.font = header_font
            c.fill = header_fill
            c.alignment = Alignment(horizontal="left", vertical="center", indent=1)
            c.border = border
        ws.row_dimensions[row].height = 22
        row += 1

        for _, en_text, es_text in items:
            ws.cell(row=row, column=1, value=en_text).alignment = wrap
            ws.cell(row=row, column=2, value=es_text).alignment = wrap
            ws.cell(row=row, column=3, value="").alignment = center
            ws.cell(row=row, column=4, value="").alignment = wrap
            for col in range(1, 5):
                ws.cell(row=row, column=col).border = border
            row += 1

        row += 1  # breathing room between sections

    # column widths: English / Spanish wide, Status narrow, Notes medium
    widths = {1: 64, 2: 64, 3: 14, 4: 36}
    for col, w in widths.items():
        ws.column_dimensions[get_column_letter(col)].width = w

    # dropdown for Status column: ✓ OK / Revise / -
    dv = DataValidation(
        type="list",
        formula1='"✓ OK,Revise,-"',
        allow_blank=True,
        showDropDown=False,  # False means show the dropdown arrow
    )
    dv.error = "Pick from the list"
    dv.errorTitle = "Invalid"
    dv.prompt = "✓ OK / Revise / -"
    dv.promptTitle = "Status"
    ws.add_data_validation(dv)
    dv.add(f"C4:C{row}")

    # Freeze title + intro so they stay visible as you scroll
    ws.freeze_panes = "A4"

    path = OUT / "butterfly-challenge-spanish.xlsx"
    wb.save(path)
    return path


# ---------- Excel summary — only the main details ----------
# Curated list of the highest-impact strings a reviewer should check first.
# Ordered the way a visitor experiences the site: top of page → bottom.
SUMMARY = [
    ("Navigation", [
        "nav.story", "nav.science", "nav.alliance", "nav.live",
        "nav.takeTheChallenge", "nav.getSupport", "nav.signIn",
    ]),
    ("Home — Hero", [
        "home.label", "home.headline", "home.subheadPart1", "home.subheadPart2",
        "home.joinCTA", "home.iDidIt",
    ]),
    ("Home — Section Headings", [
        "home.highlightsTitle", "home.communityTitle", "home.communitySub",
        "home.signTitle", "home.signSub",
        "home.howTitle", "home.effectTitle",
        "home.ctaHeadline", "home.ctaSub", "home.remindMe",
    ]),
    ("Highlights Cards", [
        "highlights.card1.title", "highlights.card2.title",
        "highlights.card3.title", "highlights.card4.title",
    ]),
    ("How It Works — Step Titles", [
        "stepTabs.step1.title", "stepTabs.step1.heading",
        "stepTabs.step2.title", "stepTabs.step2.heading",
        "stepTabs.step3.title", "stepTabs.step3.heading",
    ]),
    ("Sign Builder", [
        "signBuilder.intro.title", "signBuilder.intro.sub", "signBuilder.intro.start",
        "signBuilder.done.label", "signBuilder.done.line1",
        "signBuilder.done.line2", "signBuilder.done.line3",
    ]),
    ("Butterfly Effect Chain", [
        "chain.done.title", "chain.done.cta",
    ]),
    ("Live Feed", [
        "live.title", "live.topCountries", "live.topCities",
        "live.topParticipants", "live.shareMoment",
    ]),
    ("FAQ — Questions", [
        "faq.what.q", "faq.why.q", "faq.donate.q",
        "faq.therapy.q", "faq.month.q", "faq.who.q",
    ]),
    ("Role Cards (Alliance Page)", [
        "popups.role.creators.name", "popups.role.celebrities.name",
        "popups.role.athletes.name", "popups.role.music.name",
        "popups.role.dance.name", "popups.role.film.name",
        "popups.role.fashion.name", "popups.role.art.name",
        "popups.role.faith.name", "popups.role.gaming.name",
        "popups.role.podcast.name", "popups.role.everyone.name",
    ]),
    ("Alliance Partners", [
        "popups.alliance.platforms.name", "popups.alliance.care.name",
        "popups.alliance.media.name", "popups.alliance.business.name",
        "popups.alliance.education.name", "popups.alliance.connectivity.name",
    ]),
    ("Story Page", [
        "story.heroTitle", "story.heroSub",
        "story.cta.title", "story.cta.sub",
    ]),
    ("Science Page", [
        "science.heroTitle", "science.heroSub",
        "science.cta.title", "science.cta.sub",
    ]),
    ("Alliance Page", [
        "alliancePage.heroTitle", "alliancePage.heroSub",
        "alliancePage.ctaTitle", "alliancePage.ctaSub", "alliancePage.ctaBtn",
    ]),
    ("Live Page", [
        "livePage.heroTitle", "livePage.heroSub",
        "livePage.ctaTitle", "livePage.ctaSub",
    ]),
    ("Footer", [
        "footer.manyFlagsTitle", "footer.manyFlagsSub",
        "footer.tagline", "footer.getSupportNow", "footer.copyright",
    ]),
    ("Holding Page", [
        "workingProgress.label", "workingProgress.heading1",
        "workingProgress.heading2", "workingProgress.typed",
    ]),
]


def resolve_key(locale, dotted):
    cur = locale
    for part in dotted.split("."):
        if isinstance(cur, dict) and part in cur:
            cur = cur[part]
        else:
            return ""
    return str(cur) if not isinstance(cur, (dict, list)) else ""


def build_summary_xlsx():
    wb = Workbook()
    ws = wb.active
    ws.title = "Summary"

    title_font = Font(bold=True, size=16, color="FFFFFF")
    title_fill = PatternFill("solid", fgColor="0A7B77")
    intro_font = Font(italic=True, size=10, color="6E6E73")
    section_font = Font(bold=True, size=12, color="FFFFFF")
    section_fill = PatternFill("solid", fgColor="0A7B77")
    header_font = Font(bold=True, size=11, color="1D1D1F")
    header_fill = PatternFill("solid", fgColor="E8F5F3")
    thin = Side(border_style="thin", color="D2D2D7")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    wrap = Alignment(wrap_text=True, vertical="top", horizontal="left")
    center = Alignment(wrap_text=True, vertical="center", horizontal="center")

    ws.merge_cells("A1:D1")
    ws["A1"] = "Butterfly Challenge — Spanish at a Glance"
    ws["A1"].font = title_font
    ws["A1"].fill = title_fill
    ws["A1"].alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[1].height = 36

    ws.merge_cells("A2:D2")
    ws["A2"] = "The main strings only — what a visitor sees first on every page."
    ws["A2"].font = intro_font
    ws["A2"].alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[2].height = 20

    row = 4

    for section_title, keys in SUMMARY:
        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=4)
        c = ws.cell(row=row, column=1, value=section_title)
        c.font = section_font
        c.fill = section_fill
        c.alignment = Alignment(horizontal="left", vertical="center", indent=1)
        ws.row_dimensions[row].height = 26
        row += 1

        for col, h in enumerate(["English", "Español", "Status", "Notes"], start=1):
            cc = ws.cell(row=row, column=col, value=h)
            cc.font = header_font
            cc.fill = header_fill
            cc.alignment = Alignment(horizontal="left", vertical="center", indent=1)
            cc.border = border
        ws.row_dimensions[row].height = 22
        row += 1

        for dotted in keys:
            en_val = resolve_key(EN, dotted)
            es_val = resolve_key(ES, dotted)
            ws.cell(row=row, column=1, value=en_val).alignment = wrap
            ws.cell(row=row, column=2, value=es_val).alignment = wrap
            ws.cell(row=row, column=3, value="").alignment = center
            ws.cell(row=row, column=4, value="").alignment = wrap
            for col in range(1, 5):
                ws.cell(row=row, column=col).border = border
            row += 1

        row += 1

    widths = {1: 58, 2: 58, 3: 14, 4: 32}
    for col, w in widths.items():
        ws.column_dimensions[get_column_letter(col)].width = w

    dv = DataValidation(
        type="list", formula1='"✓ OK,Revise,-"', allow_blank=True, showDropDown=False,
    )
    dv.prompt = "✓ OK / Revise / -"
    dv.promptTitle = "Status"
    ws.add_data_validation(dv)
    dv.add(f"C4:C{row}")
    ws.freeze_panes = "A4"

    path = OUT / "butterfly-challenge-spanish-summary.xlsx"
    wb.save(path)
    return path


if __name__ == "__main__":
    docx_path = build_docx()
    xlsx_path = build_xlsx()
    summary_path = build_summary_xlsx()
    print(f"Wrote {docx_path}  ({os.path.getsize(docx_path)//1024} KB)")
    print(f"Wrote {xlsx_path}  ({os.path.getsize(xlsx_path)//1024} KB)")
    print(f"Wrote {summary_path}  ({os.path.getsize(summary_path)//1024} KB)")
